"""
Export all docs/*.md to PDF + DOCX for phone / WhatsApp sharing.
Usage:  python scripts/export_docs.py
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from markdown import markdown
from xhtml2pdf import pisa

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "shareable"  # PDFs + DOCXs live here (easy to zip/send)


def md_files() -> list[Path]:
    return sorted(DOCS.glob("*.md"))


def md_to_html(text: str, title: str) -> str:
    body = markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )
    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8"/>
<title>{title}</title>
<style>
  @page {{ size: A4; margin: 1.6cm 1.5cm; }}
  body {{
    font-family: Helvetica, Arial, sans-serif;
    font-size: 10.5pt;
    line-height: 1.45;
    color: #1a1a1a;
  }}
  h1 {{ font-size: 18pt; margin: 0 0 10pt 0; color: #0f172a; }}
  h2 {{ font-size: 13.5pt; margin: 16pt 0 6pt 0; color: #0f172a;
       border-bottom: 1px solid #cbd5e1; padding-bottom: 3pt; }}
  h3 {{ font-size: 11.5pt; margin: 12pt 0 4pt 0; color: #1e293b; }}
  h4 {{ font-size: 10.5pt; margin: 10pt 0 3pt 0; }}
  p {{ margin: 0 0 7pt 0; }}
  ul, ol {{ margin: 0 0 8pt 18pt; }}
  li {{ margin-bottom: 3pt; }}
  code {{ font-family: Courier, monospace; font-size: 9pt;
         background: #f1f5f9; padding: 1pt 3pt; }}
  pre {{ background: #f8fafc; border: 1px solid #e2e8f0; padding: 8pt;
        font-size: 8.5pt; white-space: pre-wrap; }}
  table {{ border-collapse: collapse; width: 100%; margin: 8pt 0 12pt 0;
          font-size: 9pt; }}
  th, td {{ border: 1px solid #cbd5e1; padding: 4pt 6pt; vertical-align: top; }}
  th {{ background: #f1f5f9; text-align: left; }}
  blockquote {{ margin: 8pt 0; padding: 6pt 10pt; border-left: 3pt solid #0ea5e9;
               background: #f0f9ff; color: #0c4a6e; }}
  hr {{ border: none; border-top: 1px solid #e2e8f0; margin: 14pt 0; }}
  a {{ color: #0369a1; text-decoration: none; }}
  .cover {{ margin-bottom: 14pt; }}
  .meta {{ color: #64748b; font-size: 9pt; margin-bottom: 12pt; }}
</style>
</head>
<body>
<div class="cover">
  <div class="meta">Signal · Thirdbase · Shareable export</div>
</div>
{body}
</body>
</html>"""


def write_pdf(md_path: Path, out_path: Path) -> None:
    title = md_path.stem.replace("_", " ")
    html = md_to_html(md_path.read_text(encoding="utf-8"), title)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("wb") as f:
        status = pisa.CreatePDF(html, dest=f, encoding="utf-8")
    if status.err:
        raise RuntimeError(f"PDF failed for {md_path.name}: {status.err}")


def _set_run_font(run, size=10.5, bold=False, italic=False, code=False):
    run.font.size = Pt(9 if code else size)
    run.bold = bold
    run.italic = italic
    run.font.name = "Consolas" if code else "Calibri"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Consolas" if code else "Calibri")
    rFonts.set(qn("w:hAnsi"), "Consolas" if code else "Calibri")


def _add_inline(paragraph, text: str):
    """Very small inline parser: **bold**, `code`, *italic*."""
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            _set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, code=True)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            _set_run_font(run, italic=True)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run)


def write_docx(md_path: Path, out_path: Path) -> None:
    """Lightweight MD → DOCX (headings, lists, tables, quotes, paragraphs)."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    meta = doc.add_paragraph()
    run = meta.add_run("Signal · Thirdbase · Shareable export")
    _set_run_font(run, size=9, italic=True)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                _set_run_font(run, code=True, size=8.5)
                in_code = False
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("—" * 24)
            i += 1
            continue

        # Table block
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*[-:| ]+\|", lines[i + 1] or ""):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                raw = lines[i].strip()
                if re.match(r"^\|?\s*[-:| ]+\|?$", raw):
                    i += 1
                    continue
                cells = [c.strip() for c in raw.strip("|").split("|")]
                rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = row[c_idx] if c_idx < len(row) else ""
                        for p in cell.paragraphs:
                            for run in p.runs:
                                _set_run_font(run, size=9, bold=(r_idx == 0))
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_heading(level=min(level, 3))
            p.clear()
            run = p.add_run(text)
            sizes = {1: 18, 2: 14, 3: 12, 4: 11}
            _set_run_font(run, size=sizes.get(level, 11), bold=True)
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, text)
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, text)
            i += 1
            continue

        if line.strip().startswith(">"):
            text = re.sub(r"^\s*>\s?", "", line).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            run = p.add_run(text)
            _set_run_font(run, italic=True, size=10)
            run.font.color.rgb = RGBColor(0x0C, 0x4A, 0x6E)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        _add_inline(p, line.strip())
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main() -> int:
    files = md_files()
    if not files:
        print("No markdown files in docs/", file=sys.stderr)
        return 1

    OUT.mkdir(parents=True, exist_ok=True)
    print(f"Exporting {len(files)} docs -> {OUT}")

    for md in files:
        pdf = OUT / f"{md.stem}.pdf"
        docx_path = OUT / f"{md.stem}.docx"
        print(f"  {md.name} -> PDF + DOCX")
        write_pdf(md, pdf)
        write_docx(md, docx_path)

    # Also copy into docs/ root for quick access (same folder as .md)
    for md in files:
        for ext in (".pdf", ".docx"):
            src = OUT / f"{md.stem}{ext}"
            dst = DOCS / f"{md.stem}{ext}"
            dst.write_bytes(src.read_bytes())

    print("Done.")
    print(f"  Folder: {OUT}")
    print(f"  Also mirrored next to each .md in {DOCS}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
